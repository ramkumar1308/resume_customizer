import io, json, zipfile, re
import streamlit as st
from docx import Document
from docx.shared import Pt

with open("master_resume_modules.json") as f:
    LIB = json.load(f)

def add_heading(doc,text,size=12):
    p=doc.add_paragraph()
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size)

def add_bullets(doc,items):
    for it in items:
        p=doc.add_paragraph()
        r=p.add_run("• "+it); r.font.size=Pt(10.5)

def build_resume(company,role,jd):
    doc=Document()
    add_heading(doc,LIB["header"]["name"],16)
    doc.add_paragraph(LIB["header"]["title"])
    doc.add_paragraph(LIB["header"]["contacts"])
    add_heading(doc,"SUMMARY")
    for l in LIB["summary"]: doc.add_paragraph(l)
    add_heading(doc,"EXPERIENCE")
    for exp in LIB["experience"]:
        add_heading(doc,f"{exp['company']} | {exp['role']} ({exp['dates']})",11)
        add_bullets(doc,[b['text'] for b in exp['bullets'][:3]])
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf

def build_cover(company,role,jd):
    doc=Document()
    doc.add_paragraph("Dear Hiring Manager,")
    doc.add_paragraph(f"I am applying for {role} at {company}.")
    doc.add_paragraph("Sincerely, "+LIB["header"]["name"])
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf

st.title("SRE Resume Customizer")
company=st.text_input("Company")
role=st.text_input("Role")
jd=st.text_area("Job Description")
if st.button("Generate Apply Pack"):
    if company and role and jd.strip():
        res=build_resume(company,role,jd)
        cov=build_cover(company,role,jd)
        zbuf=io.BytesIO()
        with zipfile.ZipFile(zbuf,"w",zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(f"{company}_{role}_Resume.docx",res.getvalue())
            zf.writestr(f"{company}_{role}_Cover_Letter.docx",cov.getvalue())
        zbuf.seek(0)
        st.download_button("Download",zbuf,f"{company}_{role}_Apply_Pack.zip","application/zip")
    else:
        st.error("Fill all fields")
