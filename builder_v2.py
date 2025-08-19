
import io, re
from typing import Dict, List, Set
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING

KEYWORD_TO_TAG = {
    r"\bpython\b": "python",
    r"\bflask\b": "flask",
    r"\bterraform\b": "terraform",
    r"\bansible\b": "ansible",
    r"\bkubernetes\b|\beks\b|\baks\b|\bgke\b": "kubernetes",
    r"\baws\b|\bazure\b|\bgcp\b": "cloud",
    r"\bprometheus\b|\bgrafana\b|\bjaeger\b|\bloki\b|\bsplunk\b|\bappd(ynamics)?\b": "observability",
    r"\bhigh availability\b|\bha\b|\bfailover\b|\bdisaster recovery\b|\bdr\b|\bactive-active\b": "systemdesign",
    r"\bsli(s|/)?\bslo(s)?\b|\bsli\b|\bslo\b": "slos",
    r"\bincident\b|\boncall\b|\bpostmortem\b": "incident",
    r"\bci/cd\b|\bjenkins\b|\bgitlab\b|\bazure devops\b": "cicd",
    r"\bsalesforce\b": "salesforce",
    r"\bment(or|oring|ored)\b|\blead\b|\bleadership\b": "leadership",
    r"\bnfr\b|\bnon-functional\b": "nfr",
    r"\bcapacity\b|\bscalability\b": "capacity",
}

def extract_tags(text: str) -> Set[str]:
    tags = set()
    t = text.lower()
    for pat, tag in KEYWORD_TO_TAG.items():
        if re.search(pat, t):
            tags.add(tag)
    return tags

def rank_bullets(bullets: List[Dict], wanted: Set[str], k: int) -> List[str]:
    ranked = sorted(bullets, key=lambda b: len(set(b.get("tags", [])) & wanted), reverse=True)
    out, seen = [], set()
    for b in ranked:
        t = b.get("text","").strip()
        if not t or t in seen: 
            continue
        out.append(t); seen.add(t)
        if len(out) >= k: break
    if not out:
        out = [b.get("text","") for b in bullets[:k]]
    return out

def set_paragraph_style(p, size=10.5, space_after=0.6):
    p_format = p.paragraph_format
    p_format.space_after = Pt(space_after)
    p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for r in p.runs:
        r.font.size = Pt(size)
        r.font.name = "Calibri"

def add_heading(doc, text, size=12, top=12, bottom=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(top)
    p.paragraph_format.space_after = Pt(bottom)

def add_bullets(doc, items, size=10.5):
    for t in items:
        p = doc.add_paragraph(style=None)
        run = p.add_run("• " + t)
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(1)

def build_resume_docx(lib: Dict, company: str, role: str, jd_text: str) -> io.BytesIO:
    wanted = extract_tags(jd_text)
    exp = lib.get("experience", [])
    exp_recent = exp[:2]
    exp_earlier = exp[2:]

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.6)
        s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.7)
        s.right_margin = Inches(0.7)

    hdr = lib.get("header", {})
    add_heading(doc, hdr.get("name","YOUR NAME"), size=16, top=0, bottom=2)
    p = doc.add_paragraph(hdr.get("title","")); set_paragraph_style(p, size=11, space_after=1)
    p = doc.add_paragraph(hdr.get("contacts","")); set_paragraph_style(p, size=10, space_after=2)

    add_heading(doc, "PROFESSIONAL SUMMARY", size=12, top=10, bottom=4)
    for line in lib.get("summary", []):
        p = doc.add_paragraph(line); set_paragraph_style(p, size=10.5, space_after=1)

    if lib.get("skills_groups"):
        add_heading(doc, "CORE SKILLS", size=12, top=10, bottom=4)
        for title, items in lib["skills_groups"]:
            p = doc.add_paragraph()
            r = p.add_run(f"{title}: "); r.bold = True; r.font.size = Pt(10.5); r.font.name = "Calibri"
            r2 = p.add_run(", ".join(items)); r2.font.size = Pt(10.5); r2.font.name = "Calibri"
            set_paragraph_style(p, size=10.5, space_after=0.5)

    add_heading(doc, "PROFESSIONAL EXPERIENCE", size=12, top=10, bottom=4)
    for ex in exp_recent:
        add_heading(doc, f"{ex.get('company','')} | {ex.get('role','')} ({ex.get('dates','')})", size=11, top=8, bottom=2)
        add_bullets(doc, rank_bullets(ex.get("bullets", []), wanted, 6), size=10.5)

    doc.add_page_break()

    for ex in exp_earlier:
        add_heading(doc, f"{ex.get('company','')} | {ex.get('role','')} ({ex.get('dates','')})", size=11, top=8, bottom=2)
        add_bullets(doc, rank_bullets(ex.get("bullets", []), wanted, 4), size=10.5)

    if lib.get("projects"):
        add_heading(doc, "KEY PROJECTS", size=12, top=10, bottom=3)
        # pick top 4 by match
        ranked = sorted(lib["projects"], key=lambda p: len(set(p.get("tags", [])) & wanted), reverse=True)
        add_bullets(doc, [p["text"] for p in ranked[:4]], size=10.5)

    if lib.get("certs"):
        add_heading(doc, "CERTIFICATIONS", size=12, top=10, bottom=3)
        add_bullets(doc, lib["certs"], size=10.5)

    if lib.get("education"):
        add_heading(doc, "EDUCATION", size=12, top=10, bottom=3)
        add_bullets(doc, lib["education"], size=10.5)

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf

def build_cover_letter_docx(lib: Dict, company: str, role: str, jd_text: str) -> io.BytesIO:
    tags = extract_tags(jd_text)
    strengths = []
    if "python" in tags: strengths.append("Python automation and tooling")
    if "systemdesign" in tags: strengths.append("high availability & fault tolerance")
    if "observability" in tags: strengths.append("metrics, logs, tracing")
    if "kubernetes" in tags: strengths.append("Kubernetes at scale")
    if "cicd" in tags: strengths.append("modern CI/CD")
    if "slos" in tags: strengths.append("SLIs/SLOs & incident reduction")
    if not strengths: strengths.append("site reliability and cloud operations")

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.7)
        s.right_margin = Inches(0.7)

    hdr = lib.get("header", {})
    p = doc.add_paragraph(hdr.get("name","") + " | " + hdr.get("contacts",""))
    for r in p.runs: r.font.size = Pt(10); r.font.name = "Calibri"
    doc.add_paragraph("")
    pr = doc.add_paragraph("Dear Hiring Manager,"); 
    for r in pr.runs: r.font.size = Pt(11); r.font.name = "Calibri"
    p = doc.add_paragraph(f"I’m excited to apply for the {role} role at {company}. I bring 14+ years in SRE with strengths in {', '.join(strengths)}.")
    for r in p.runs: r.font.size = Pt(10.5); r.font.name = "Calibri"
    for b in [
        "Directed SRE strategy and built Python automation saving 20+ hours/week (Wells Fargo).",
        "Developed a Python/Flask risk tool integrated with Gremlin for chaos planning (Morgan Stanley).",
        "Built a Self‑Service NFR Logging Portal to standardize logging quality across services.",
        "Implemented distributed tracing and synthetic monitoring to improve early detection by 35%."
    ]:
        p = doc.add_paragraph("• " + b)
        for r in p.runs: r.font.size = Pt(10.5); r.font.name = "Calibri"
    p = doc.add_paragraph("I’d welcome the chance to discuss how I can strengthen reliability and engineering velocity for your team.")
    for r in p.runs: r.font.size = Pt(10.5); r.font.name = "Calibri"
    doc.add_paragraph("Thank you for your time and consideration.")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(hdr.get("name",""))

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf
