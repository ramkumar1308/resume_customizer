
import io, re, json
from typing import Dict, List, Set
from docx import Document
from docx.shared import Pt, Inches

KEYWORD_TO_TAG = {
    r"\bpython\b": "python",
    r"\bflask\b": "flask",
    r"\bterraform\b": "terraform",
    r"\bansible\b": "ansible",
    r"\bkubernetes\b|\beks\b|\baks\b|\bgke\b": "kubernetes",
    r"\baws\b|\bazure\b|\bgcp\b": "cloud",
    r"\bprometheus\b|\bgrafana\b|\bjaeger\b|\bloki\b|\bsplunk\b|\bappd(ynamics)?\b": "observability",
    r"\bhigh availability\b|\bha\b|\bfailover\b|\bdisaster recovery\b|\bdr\b|\bactive-active\b": "systemdesign",
    r"\bsli(s|/)?\bslo(s)?\b|\bsli\b|\bslo\b": "slos",
    r"\bincident\b|\boncall\b|\bpostmortem\b": "incident",
    r"\bci/cd\b|\bjenkins\b|\bgitlab\b|\bazure devops\b": "cicd",
    r"\bsalesforce\b": "salesforce",
    r"\bment(or|oring|ored)\b|\blead\b|\bleadership\b": "leadership",
    r"\bnfr\b|\bnon-functional\b": "nfr",
    r"\bcapacity\b|\bscalability\b": "capacity",
}

def extract_tags_from_jd(jd_text: str) -> Set[str]:
    tags = set()
    text = jd_text.lower()
    for pattern, tag in KEYWORD_TO_TAG.items():
        if re.search(pattern, text):
            tags.add(tag)
    return tags

def pick_bullets(bullets: List[Dict], wanted_tags: Set[str], k: int) -> List[str]:
    ranked = sorted(bullets, key=lambda b: len(set(b.get("tags", [])) & wanted_tags), reverse=True)
    out, seen = [], set()
    for b in ranked:
        t = b["text"]
        if t not in seen:
            out.append(t)
            seen.add(t)
        if len(out) >= k:
            break
    if not out:
        out = [b["text"] for b in bullets[:k]]
    return out

def add_heading(doc, text, size=12, bold=True, space_before=12, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)

def add_bullets(doc, items, font_size=10.5, space_after=1):
    for it in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(f"• {it}")
        run.font.size = Pt(font_size)

def build_resume_docx(lib: Dict, company: str, role: str, jd_text: str) -> io.BytesIO:
    wanted_tags = extract_tags_from_jd(jd_text)
    exp = lib.get("experience", [])
    exp_recent = exp[:2]
    exp_earlier = exp[2:]

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.7)
        s.right_margin = Inches(0.7)

    hdr = lib.get("header", {})
    add_heading(doc, hdr.get("name", "YOUR NAME"), size=16, space_before=0, space_after=2)
    p = doc.add_paragraph(hdr.get("title", "")); 
    if p.runs: p.runs[0].font.size = Pt(11)
    p = doc.add_paragraph(hdr.get("contacts", "")); 
    if p.runs: p.runs[0].font.size = Pt(10)

    add_heading(doc, "PROFESSIONAL SUMMARY", size=12)
    for line in lib.get("summary", []):
        pr = doc.add_paragraph(line); 
        if pr.runs: pr.runs[0].font.size = Pt(10.5)

    if lib.get("skills_groups"):
        add_heading(doc, "CORE SKILLS", size=12)
        for title, items in lib["skills_groups"]:
            p = doc.add_paragraph()
            r = p.add_run(f"{title}: "); r.bold = True; r.font.size = Pt(10.5)
            r2 = p.add_run(", ".join(items)); r2.font.size = Pt(10.5)

    add_heading(doc, "PROFESSIONAL EXPERIENCE", size=12)

    for ex in exp_recent:
        add_heading(doc, f"{ex.get('company','')} | {ex.get('role','')} ({ex.get('dates','')})", size=11, space_before=8, space_after=4)
        add_bullets(doc, pick_bullets(ex.get("bullets", []), wanted_tags, 6))

    doc.add_page_break()

    for ex in exp_earlier:
        add_heading(doc, f"{ex.get('company','')} | {ex.get('role','')} ({ex.get('dates','')})", size=11, space_before=8, space_after=4)
        add_bullets(doc, pick_bullets(ex.get("bullets", []), wanted_tags, 4))

    if lib.get("projects"):
        add_heading(doc, "KEY PROJECTS", size=12, space_before=12)
        ranked = sorted(lib["projects"], key=lambda p: len(set(p.get("tags", [])) & wanted_tags), reverse=True)
        add_bullets(doc, [p["text"] for p in ranked[:4]])

    if lib.get("certs"):
        add_heading(doc, "CERTIFICATIONS", size=12, space_before=12)
        add_bullets(doc, lib["certs"], font_size=10.5)

    if lib.get("education"):
        add_heading(doc, "EDUCATION", size=12, space_before=12)
        add_bullets(doc, lib["education"], font_size=10.5)

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf

def build_cover_letter_docx(lib: Dict, company: str, role: str, jd_text: str) -> io.BytesIO:
    tags = extract_tags_from_jd(jd_text)
    strengths = []
    if "python" in tags: strengths.append("automation and tooling in Python")
    if "systemdesign" in tags: strengths.append("designing highly available, fault‑tolerant systems")
    if "observability" in tags: strengths.append("building observability frameworks (metrics, logs, tracing)")
    if "kubernetes" in tags: strengths.append("operating Kubernetes at scale")
    if "cicd" in tags: strengths.append("modernizing CI/CD pipelines")
    if "slos" in tags: strengths.append("defining SLIs/SLOs and improving incident response")
    if not strengths: strengths.append("site reliability, automation, and cloud operations")

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.7)
        s.right_margin = Inches(0.7)

    hdr = lib.get("header", {})
    p = doc.add_paragraph(hdr.get("name","") + " | " + hdr.get("contacts",""))
    if p.runs: p.runs[0].font.size = Pt(10)
    doc.add_paragraph("")
    pr = doc.add_paragraph("Dear Hiring Manager,"); 
    if pr.runs: pr.runs[0].font.size = Pt(11)
    doc.add_paragraph(
        f"I’m excited to apply for the {role} role at {company}. I bring 14+ years in enterprise SRE with strengths in {', '.join(strengths)}."
    )
    for b in [
        "Directed SRE strategy and built Python automation saving 20+ hours/week (Wells Fargo).",
        "Developed a Python/Flask risk tool integrated with Gremlin for chaos planning (Morgan Stanley).",
        "Built a Self‑Service NFR Logging Portal to standardize logging quality across services.",
        "Implemented distributed tracing and synthetic monitoring to improve early detection by 35%."
    ]:
        p = doc.add_paragraph(); r = p.add_run("• " + b); r.font.size = Pt(10.5)
    doc.add_paragraph("I’d welcome the chance to discuss how I can strengthen reliability and engineering velocity for your team.")
    doc.add_paragraph("Thank you for your time and consideration.")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(hdr.get("name",""))

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf
